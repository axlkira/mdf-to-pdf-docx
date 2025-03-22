import sys
import os
from PySide6.QtWidgets import (QApplication, QMainWindow, QPushButton, QVBoxLayout, 
                            QWidget, QFileDialog, QLabel, QMessageBox, QHBoxLayout)
from PySide6.QtCore import Qt
from PySide6.QtGui import QPalette, QColor
import markdown2
import tempfile
import pdfkit
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MarkdownToPDFConverter(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()
        
    def initUI(self):
        # Configuración de la ventana principal
        self.setWindowTitle('Markdown a PDF Converter')
        self.setGeometry(100, 100, 600, 400)
        
        # Widget central y layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # Estilo para los widgets
        self.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px;
                border-radius: 5px;
                font-size: 14px;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QLabel {
                font-size: 14px;
                color: #2c3e50;
                margin: 10px;
            }
        """)
        
        # Etiqueta de título
        title_label = QLabel('Conversor de Markdown a PDF')
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet('font-size: 24px; font-weight: bold; margin: 20px;')
        layout.addWidget(title_label)
        
        # Etiqueta de instrucciones
        instructions = QLabel('Selecciona un archivo Markdown (.md) para convertirlo a PDF o DOCX')
        instructions.setAlignment(Qt.AlignCenter)
        layout.addWidget(instructions)
        
        # Contenedor para botones
        button_container = QWidget()
        button_layout = QHBoxLayout(button_container)
        button_layout.setContentsMargins(0, 0, 0, 0)
        button_layout.setSpacing(10)
        button_layout.setAlignment(Qt.AlignCenter)
        
        # Botón para seleccionar archivo y convertir a PDF
        self.pdf_button = QPushButton('Convertir a PDF')
        self.pdf_button.clicked.connect(lambda: self.select_file_and_convert('pdf'))
        button_layout.addWidget(self.pdf_button)
        
        # Botón para seleccionar archivo y convertir a DOCX
        self.docx_button = QPushButton('Convertir a DOCX')
        self.docx_button.clicked.connect(lambda: self.select_file_and_convert('docx'))
        button_layout.addWidget(self.docx_button)
        
        layout.addWidget(button_container)
        
        # Etiqueta para mostrar el archivo seleccionado
        self.file_label = QLabel('Ningún archivo seleccionado')
        self.file_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.file_label)
        
        # Agregar espacio
        layout.addStretch()
        
        # Etiqueta de atribución
        attribution_label = QLabel('Hecho por Alexander Muñoz Castro')
        attribution_label.setAlignment(Qt.AlignCenter)
        attribution_label.setStyleSheet('font-size: 12px; font-style: italic; color: #7f8c8d;')
        layout.addWidget(attribution_label)

        # Configurar colores de fondo
        palette = self.palette()
        palette.setColor(QPalette.Window, QColor('#f5f6fa'))
        self.setPalette(palette)

    def select_file_and_convert(self, file_type):
        file_name, _ = QFileDialog.getOpenFileName(
            self,
            "Seleccionar archivo Markdown",
            "",
            "Markdown Files (*.md)"
        )
        
        if file_name:
            self.file_label.setText(f'Archivo seleccionado: {os.path.basename(file_name)}')
            if file_type == 'pdf':
                self.convert_to_pdf(file_name)
            elif file_type == 'docx':
                self.convert_to_docx(file_name)

    def convert_to_pdf(self, md_file):
        try:
            # Leer el contenido del archivo Markdown
            with open(md_file, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convertir Markdown a HTML con extras
            html_content = markdown2.markdown(
                md_content,
                extras=[
                    'tables',
                    'fenced-code-blocks',
                    'header-ids',
                    'tag-friendly',
                    'break-on-newline'
                ]
            )
            
            # Estilos CSS mejorados con más espaciado y colores
            css_content = """
                body {
                    font-family: 'Arial', sans-serif;
                    line-height: 1.8;
                    color: #2c3e50;
                    max-width: 21cm;
                    margin: 0 auto;
                    padding: 2cm;
                    background-color: #ffffff;
                }
                h1, h2, h3, h4, h5, h6 {
                    margin-top: 2em;
                    margin-bottom: 1em;
                    page-break-after: avoid;
                }
                h1 {
                    color: #2c3e50;
                    font-size: 2.5em;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 0.5em;
                }
                h2 {
                    color: #34495e;
                    font-size: 2em;
                    border-bottom: 2px solid #95a5a6;
                    padding-bottom: 0.3em;
                }
                h3 {
                    color: #2980b9;
                    font-size: 1.5em;
                }
                p {
                    margin: 1.2em 0;
                    text-align: justify;
                    line-height: 1.8;
                }
                ul, ol {
                    padding-left: 2em;
                    margin: 1em 0;
                }
                li {
                    margin: 0.5em 0;
                    line-height: 1.6;
                }
                a {
                    color: #3498db;
                    text-decoration: none;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 2em 0;
                    page-break-inside: avoid;
                }
                th {
                    background-color: #3498db;
                    color: white;
                    font-weight: bold;
                    padding: 12px;
                    border: 1px solid #bdc3c7;
                }
                td {
                    padding: 12px;
                    border: 1px solid #bdc3c7;
                }
                tr:nth-child(even) {
                    background-color: #f8f9fa;
                }
                tr:hover {
                    background-color: #edf2f7;
                }
                code {
                    background-color: #f8f9fa;
                    color: #e74c3c;
                    padding: 0.2em 0.4em;
                    border-radius: 4px;
                    font-family: 'Courier New', monospace;
                    font-size: 0.9em;
                }
                pre {
                    background-color: #2c3e50;
                    color: #ecf0f1;
                    padding: 1em;
                    border-radius: 8px;
                    overflow-x: auto;
                    margin: 1.5em 0;
                    page-break-inside: avoid;
                }
                pre code {
                    background-color: transparent;
                    color: inherit;
                    padding: 0;
                    border-radius: 0;
                }
                blockquote {
                    border-left: 4px solid #3498db;
                    margin: 1.5em 0;
                    padding: 1em 1.5em;
                    background-color: #f8f9fa;
                    color: #34495e;
                    font-style: italic;
                    page-break-inside: avoid;
                }
                img {
                    max-width: 100%;
                    height: auto;
                    display: block;
                    margin: 2em auto;
                    page-break-inside: avoid;
                }
                hr {
                    border: none;
                    border-top: 2px solid #bdc3c7;
                    margin: 2em 0;
                }
                /* Ajustes para impresión */
                @media print {
                    body {
                        font-size: 12pt;
                    }
                    h1 {
                        font-size: 24pt;
                    }
                    h2 {
                        font-size: 20pt;
                    }
                    h3 {
                        font-size: 16pt;
                    }
                }
            """
            
            # Crear HTML completo con estilos incrustados
            full_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        {css_content}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
            """
            
            # Crear archivo HTML temporal
            temp_html = tempfile.NamedTemporaryFile(delete=False, suffix='.html')
            temp_html.write(full_html.encode('utf-8'))
            temp_html.close()
            
            # Generar el nombre del archivo PDF
            pdf_file = os.path.splitext(md_file)[0] + '.pdf'
            
            # Verificar si el archivo PDF ya existe
            if os.path.exists(pdf_file):
                reply = QMessageBox.question(
                    self,
                    "Archivo existente",
                    f"El archivo {os.path.basename(pdf_file)} ya existe.\n¿Desea sobrescribirlo?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                
                if reply == QMessageBox.No:
                    QMessageBox.information(
                        self,
                        "Operación cancelada",
                        "La conversión ha sido cancelada."
                    )
                    return
            
            # Configuración de wkhtmltopdf
            config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')
            
            # Opciones mejoradas para el PDF
            options = {
                'encoding': 'UTF-8',
                'margin-top': '25mm',
                'margin-right': '25mm',
                'margin-bottom': '25mm',
                'margin-left': '25mm',
                'page-size': 'A4',
                'enable-local-file-access': None,
                'header-font-size': '9',
                'header-spacing': '10',
                'footer-right': '[page] de [topage]',
                'footer-font-size': '9',
                'footer-spacing': '10',
                'enable-smart-shrinking': True,
                'zoom': '1.2'
            }
            
            # Convertir a PDF
            pdfkit.from_file(temp_html.name, pdf_file, configuration=config, options=options)
            
            # Eliminar archivo temporal
            os.unlink(temp_html.name)
            
            QMessageBox.information(
                self,
                "Éxito",
                f"PDF generado exitosamente:\n{pdf_file}"
            )
            
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Error al convertir el archivo:\n{str(e)}"
            )

    def convert_to_docx(self, md_file):
        try:
            # Leer el contenido del archivo Markdown
            with open(md_file, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Crear documento DOCX
            doc = docx.Document()
            
            # Extraer el nombre del archivo para usarlo como título
            file_name = os.path.basename(md_file)
            file_title = os.path.splitext(file_name)[0]
            
            # Agregar título basado en el nombre del archivo
            title = doc.add_heading(file_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Procesar el contenido Markdown línea por línea
            current_list = None
            list_level = 0
            in_code_block = False
            code_block_content = []
            
            # Dividir el contenido en líneas
            lines = md_content.split('\n')
            i = 0
            
            while i < len(lines):
                try:
                    line = lines[i].strip()
                    
                    # Detectar bloques de código
                    if line.startswith('```'):
                        in_code_block = not in_code_block
                        if in_code_block:
                            code_block_content = []
                        else:
                            # Agregar el bloque de código completo
                            code_para = doc.add_paragraph()
                            code_run = code_para.add_run('\n'.join(code_block_content))
                            code_run.font.name = 'Courier New'
                            code_run.font.size = Pt(9)
                            
                            # Agregar sombreado al párrafo
                            code_para.paragraph_format.left_indent = Pt(20)
                            code_para.paragraph_format.right_indent = Pt(20)
                            code_para.paragraph_format.space_before = Pt(10)
                            code_para.paragraph_format.space_after = Pt(10)
                        i += 1
                        continue
                    
                    # Si estamos dentro de un bloque de código, agregar la línea al contenido
                    if in_code_block:
                        code_block_content.append(line)
                        i += 1
                        continue
                    
                    # Encabezados (h1 a h6)
                    if line.startswith('#'):
                        level = 0
                        for char in line:
                            if char == '#':
                                level += 1
                            else:
                                break
                        
                        if level >= 1 and level <= 6:
                            heading_text = line[level:].strip()
                            heading = doc.add_heading(heading_text, level)
                            heading.paragraph_format.space_before = Pt(12)
                            heading.paragraph_format.space_after = Pt(6)
                    
                    # Listas no ordenadas
                    elif line.startswith('* ') or line.startswith('- '):
                        item_text = line[2:].strip()
                        doc.add_paragraph(item_text, style='List Bullet')
                    
                    # Listas ordenadas
                    elif line and line[0].isdigit() and '. ' in line:
                        dot_pos = line.find('. ')
                        if dot_pos != -1:
                            item_text = line[dot_pos+2:].strip()
                            doc.add_paragraph(item_text, style='List Number')
                        else:
                            # Si no tiene el formato esperado, tratar como texto normal
                            para = doc.add_paragraph(line)
                    
                    # Citas
                    elif line.startswith('>'):
                        quote_text = line[1:].strip()
                        quote = doc.add_paragraph(quote_text)
                        quote.paragraph_format.left_indent = Pt(24)
                        quote.paragraph_format.right_indent = Pt(24)
                        quote.paragraph_format.space_before = Pt(6)
                        quote.paragraph_format.space_after = Pt(6)
                        quote.style = 'Intense Quote'
                    
                    # Líneas horizontales
                    elif line == '---' or line == '***' or line == '___':
                        doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                    
                    # Párrafos normales (texto)
                    elif line:
                        # Procesar formato inline
                        para = doc.add_paragraph()
                        
                        # Método simplificado para procesar el texto
                        # Agregar el texto sin formato especial
                        para.add_run(line)
                    
                    # Línea vacía
                    else:
                        doc.add_paragraph()
                    
                    i += 1
                except Exception as e:
                    # Si hay un error al procesar una línea, simplemente la agregamos como texto plano
                    # y continuamos con la siguiente línea
                    para = doc.add_paragraph(lines[i] if i < len(lines) else "")
                    i += 1
            
            # Generar el nombre del archivo DOCX
            docx_file = os.path.splitext(md_file)[0] + '.docx'
            
            # Verificar si el archivo DOCX ya existe
            if os.path.exists(docx_file):
                reply = QMessageBox.question(
                    self,
                    "Archivo existente",
                    f"El archivo {os.path.basename(docx_file)} ya existe.\n¿Desea sobrescribirlo?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                
                if reply == QMessageBox.No:
                    QMessageBox.information(
                        self,
                        "Operación cancelada",
                        "La conversión ha sido cancelada."
                    )
                    return
            
            # Agregar atribución al final del documento
            doc.add_paragraph()  # Espacio en blanco
            attribution = doc.add_paragraph()
            attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER
            attribution_run = attribution.add_run("Hecho por Alexander Muñoz Castro")
            attribution_run.font.size = Pt(10)
            attribution_run.font.italic = True
            attribution_run.font.color.rgb = RGBColor(127, 140, 141)  # Color gris similar al de la UI
            
            # Guardar el archivo DOCX
            doc.save(docx_file)
            
            QMessageBox.information(
                self,
                "Éxito",
                f"DOCX generado exitosamente:\n{docx_file}"
            )
            
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Error al convertir el archivo:\n{str(e)}"
            )

def main():
    app = QApplication(sys.argv)
    converter = MarkdownToPDFConverter()
    converter.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
